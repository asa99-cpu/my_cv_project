import streamlit as st
from components import display_personal_details, display_skills, display_education, display_languages, display_internships
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Page configuration
st.set_page_config(
    page_title="Derin Najmadin Mahamd - CV",
    page_icon="📄",
    layout="wide"
)

# Custom CSS
with open("styles/styles.css") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

# Header
st.markdown("""
    <div class="header">
        <h1>Derin Najmadin Mahamd</h1>
        <h3>Graphic Designer & Full-Stack Developer</h3>
    </div>
""", unsafe_allow_html=True)

# Sidebar for image
with st.sidebar:
    st.image("assets/cv_image.jpg", width=200)
    st.markdown("### Contact Information")
    st.markdown("**Email:** deman.najmadin90@gmail.com")
    st.markdown("**Phone:** +0750 710 40 32")
    st.markdown("**Date of Birth:** September 9, 1995")
    st.markdown("**Gender:** Female")
    st.markdown("**Nationality:** Kurdish")

# Main content
st.markdown("<div class='section-header'>Personal Details</div>", unsafe_allow_html=True)
display_personal_details()

st.markdown("<div class='section-header'>Skills</div>", unsafe_allow_html=True)
display_skills()

st.markdown("<div class='section-header'>Education</div>", unsafe_allow_html=True)
display_education()

st.markdown("<div class='section-header'>Languages</div>", unsafe_allow_html=True)
display_languages()

st.markdown("<div class='section-header'>Internships</div>", unsafe_allow_html=True)
display_internships()

# Function to generate PDF
def generate_pdf():
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    # Title and personal details
    c.setFont("Helvetica-Bold", 18)
    c.drawString(100, 750, "Derin Najmadin Mahamd")
    c.setFont("Helvetica", 12)
    c.drawString(100, 730, "Email: deman.najmadin90@gmail.com")
    c.drawString(100, 710, "Phone: +0750 710 40 32")
    c.drawString(100, 690, "Date of Birth: September 9, 1995")
    c.drawString(100, 670, "Gender: Female")
    c.drawString(100, 650, "Nationality: Kurdish")

    # Add skills
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, 630, "Skills:")
    c.setFont("Helvetica", 12)
    skills = ["Laboratory Technician", "Microsoft Office", "Translator", "Video Editing", "Sewing"]
    y_pos = 610
    for skill in skills:
        c.drawString(120, y_pos, f"- {skill}")
        y_pos -= 20

    # Add Education
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, y_pos, "Education:")
    y_pos -= 20
    c.setFont("Helvetica", 12)
    c.drawString(120, y_pos, "Bachelor Degree in Science of Chemistry")

    # Add Languages
    y_pos -= 20
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, y_pos, "Languages:")
    y_pos -= 20
    languages = ["Kurdish", "Arabic", "English", "Persian"]
    for language in languages:
        c.setFont("Helvetica", 12)
        c.drawString(120, y_pos, f"- {language}")
        y_pos -= 20

    # Add Internships
    y_pos -= 20
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, y_pos, "Internships:")
    y_pos -= 20
    c.setFont("Helvetica", 12)
    c.drawString(120, y_pos, "Internship Program of the Kurdistan Regional Government (2017)")

    # Save and return PDF
    c.save()
    buffer.seek(0)
    return buffer

# Function to generate Word
def generate_word():
    doc = Document()

    # Title and personal details
    title = doc.add_paragraph()
    title.add_run('Derin Najmadin Mahamd').bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Email: deman.najmadin90@gmail.com')
    doc.add_paragraph('Phone: +0750 710 40 32')
    doc.add_paragraph('Date of Birth: September 9, 1995')
    doc.add_paragraph('Gender: Female')
    doc.add_paragraph('Nationality: Kurdish')

    # Add Skills
    doc.add_paragraph('Skills:', style='Heading 1')
    skills = ["Laboratory Technician", "Microsoft Office", "Translator", "Video Editing", "Sewing"]
    for skill in skills:
        doc.add_paragraph(f"- {skill}")

    # Add Education
    doc.add_paragraph('Education:', style='Heading 1')
    doc.add_paragraph('Bachelor Degree in Science of Chemistry')

    # Add Languages
    doc.add_paragraph('Languages:', style='Heading 1')
    languages = ["Kurdish", "Arabic", "English", "Persian"]
    for language in languages:
        doc.add_paragraph(f"- {language}")

    # Add Internships
    doc.add_paragraph('Internships:', style='Heading 1')
    doc.add_paragraph('Internship Program of the Kurdistan Regional Government (2017)')

    # Save and return Word document
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Buttons to download PDF and Word
col1, col2 = st.columns(2)

with col1:
    pdf_buffer = generate_pdf()
    st.download_button(
        label="Download CV as PDF",
        data=pdf_buffer,
        file_name="Derin_Najmadin_Mahamd_CV.pdf",
        mime="application/pdf"
    )

with col2:
    word_buffer = generate_word()
    st.download_button(
        label="Download CV as Word",
        data=word_buffer,
        file_name="Derin_Najmadin_Mahamd_CV.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
